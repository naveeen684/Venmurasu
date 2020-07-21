from bs4 import BeautifulSoup
import tamil
from urllib.request import urlopen
import pypandoc
import pdfkit
import docx

# old = "https://venmurasu.in/2014/01/01/%E0%AE%B5%E0%AF%86%E0%AE%A3%E0%AF%8D%E0%AE%AE%E0%AF%81%E0%AE%B0%E0%AE%9A%E0%AF%81-%E0%AE%A8%E0%AF%82%E0%AE%B2%E0%AF%8D-%E0%AE%92%E0%AE%A9%E0%AF%8D%E0%AE%B1%E0%AF%81/"
old = "https://venmurasu.in/2015/11/23/%e0%ae%a8%e0%af%82%e0%ae%b2%e0%af%8d-%e0%ae%8e%e0%ae%9f%e0%af%8d%e0%ae%9f%e0%af%81-%e0%ae%95%e0%ae%be%e0%ae%a3%e0%af%8d%e0%ae%9f%e0%af%80%e0%ae%aa%e0%ae%ae%e0%af%8d-70/"
i = 600

# 599 https://venmurasu.in/2015/11/22/%e0%ae%a8%e0%af%82%e0%ae%b2%e0%af%8d-%e0%ae%8e%e0%ae%9f%e0%af%8d%e0%ae%9f%e0%af%81-%e0%ae%95%e0%ae%be%e0%ae%a3%e0%af%8d%e0%ae%9f%e0%af%80%e0%ae%aa%e0%ae%ae%e0%af%8d-69/

while old:

    soup = BeautifulSoup(urlopen(old), "html.parser")
    header = soup.find('header', {'class': 'entry-header'}).getText()
    content = soup.find('div', {'class': 'entry-content'}).text
    content = content.split('Share')[0]
    # print(header)
    # print(content)

    body = soup.body.text
    words = []
    letters = tamil.utf8.get_letters(body)
    word = tamil.utf8.get_tamil_words(letters)

    for pos, word in enumerate(word):
        words.append(word)

    print(i, old)

    year = int(old[21:25])
    month = int(old[26:28])
    date = int(old[29:31])
    print(year, month, date)

    mydoc = docx.Document()
    mydoc.add_heading(str(header), 1)
    mydoc.add_paragraph(str(content))

    # save file
    if i < 995:
        with open("data1/%s-%s-%s.txt" % (date, month, year), "w", encoding="utf-8") as f:
            for s in words:
                f.write(str(s) + "\n")

        mydoc.save("blog1/%s-%s-%s.rtf" % (date, month, year))

    else:
        with open("data2/%s-%s-%s.txt" % (date, month, year), "w", encoding="utf-8") as f:
            for s in words:
                f.write(str(s) + "\n")

        mydoc.save("blog2/%s-%s-%s.rtf" % (date, month, year))

    i += 1

    # check whether the next blog is available or not
    try:
        new = soup.find('div', {'class': 'nav-next'})
        new = new.find('a')
        new = new.get('href')
        old = new
    except Exception as e:
        old = False
